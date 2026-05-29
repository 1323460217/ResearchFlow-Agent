import logging
from pathlib import Path
from typing import List

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from backend.core.config import settings
from backend.core.exceptions import DocumentParseError
from backend.models.document import Document
from backend.models.document_chunk import DocumentChunk
from backend.models.knowledge_base import KnowledgeBase
from backend.rag.chroma_client import (
    delete_collection,
    get_or_create_collection,
)
from backend.rag.chunker import Chunk, chunk_document
from backend.rag.embeddings import embed_texts

logger = logging.getLogger(__name__)


# ── 文档解析 ────────────────────────────────────────────


async def _parse_pdf(file_path: str) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(file_path)
    try:
        parts = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                # Strip null bytes that break PostgreSQL UTF-8 validation
                text = text.replace("\x00", "")
                parts.append(text)

            # 提取表格（从 PDF 注释中检测）
            tables = page.find_tables()
            for table in tables:
                rows = table.extract()
                if rows:
                    md_table = _rows_to_markdown_table(rows)
                    parts.append(md_table)
        return "\n\n".join(parts)
    finally:
        doc.close()


async def _parse_docx(file_path: str) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
            parts.append(f"{'#' * level} {para.text}")
        elif para.text.strip():
            parts.append(para.text)

    # 提取表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            parts.append(_rows_to_markdown_table(rows))

    return "\n\n".join(parts)


async def _parse_html(file_path: str) -> str:
    from bs4 import BeautifulSoup
    with open(file_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return soup.get_text("\n\n", strip=True)


async def _parse_text(file_path: str) -> str:
    content = Path(file_path).read_text(encoding="utf-8")
    return content


def _rows_to_markdown_table(rows: List[List[str]]) -> str:
    if not rows:
        return ""
    lines = []
    # header
    lines.append("| " + " | ".join(str(c) for c in rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    for row in rows[1:]:
        padded = list(row) + [""] * (len(rows[0]) - len(row))
        lines.append("| " + " | ".join(str(c) for c in padded[:len(rows[0])]) + " |")
    return "\n".join(lines)


PARSERS = {
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "html": _parse_html,
    "md": _parse_text,
    "txt": _parse_text,
}


# ── 入库管线 ────────────────────────────────────────────


async def ingest_document(
    document_id: int,
    db: AsyncSession,
    chunk_strategy: str = "auto",
) -> None:
    """完整入库管线：解析 → 切片 → 向量化 → ChromaDB → PostgreSQL。

    同步执行，在 API 请求上下文中调用。
    """
    # 1. 加载 Document 记录
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise DocumentParseError(detail=f"文档 {document_id} 不存在")

    # 2. 加载 KnowledgeBase（获取 collection_name）
    kb_result = await db.execute(
        select(KnowledgeBase).where(KnowledgeBase.id == doc.knowledge_base_id)
    )
    kb = kb_result.scalar_one_or_none()
    if not kb:
        raise DocumentParseError(detail=f"知识库 {doc.knowledge_base_id} 不存在")

    collection_name = kb.collection_name  # kb_{id}

    try:
        # 标记为处理中
        doc.ingestion_status = "processing"
        doc.ingestion_error = None
        await db.flush()

        # 3. 解析文档
        file_path = doc.file_path
        if not file_path or not Path(file_path).exists():
            raise DocumentParseError(detail=f"文件不存在: {file_path}")

        parser = PARSERS.get(doc.file_type, _parse_text)
        logger.info("Parsing document %d (%s)", document_id, doc.file_type)
        full_text = await parser(file_path)

        if not full_text.strip():
            raise DocumentParseError(detail="文档内容为空")

        # 4. 切片
        chunks: List[Chunk] = chunk_document(full_text, strategy=chunk_strategy)
        if not chunks:
            raise DocumentParseError(detail="文档切片结果为空")

        logger.info("Document %d chunked into %d pieces", document_id, len(chunks))

        # 5. 向量化
        texts = [c.content for c in chunks]
        embeddings = await embed_texts(texts)

        # 6. 写入 ChromaDB
        collection = get_or_create_collection(collection_name)
        chroma_ids = [f"doc_{document_id}_chunk_{c.chunk_index}" for c in chunks]

        collection.add(
            ids=chroma_ids,
            embeddings=embeddings,
            documents=texts,
            metadatas=[
                {
                    "document_id": document_id,
                    "chunk_index": c.chunk_index,
                    "kb_id": doc.knowledge_base_id,
                    "source": doc.filename,
                    **c.metadata,
                }
                for c in chunks
            ],
        )

        # 7. 写入 PostgreSQL（chunk 记录）
        for idx, (chunk, chroma_id) in enumerate(zip(chunks, chroma_ids)):
            chunk_record = DocumentChunk(
                document_id=document_id,
                chunk_index=chunk.chunk_index,
                content=chunk.content,
                token_count=chunk.token_count,
                chroma_id=chroma_id,
                extra_metadata=chunk.metadata,
            )
            db.add(chunk_record)

        # 8. 更新 Document 状态
        doc.ingestion_status = "done"
        doc.ingestion_error = None
        kb.chunk_count += len(chunks)
        await db.flush()

        from backend.rag.retrieval import invalidate_cache

        invalidate_cache(collection_name)
        logger.info(
            "Document %d ingestion complete: %d chunks stored in collection %s",
            document_id,
            len(chunks),
            collection_name,
        )

    except Exception as exc:
        # 回滚 ChromaDB 中已写入的 chunks（如果有）
        try:
            chroma_ids = [
                f"doc_{document_id}_chunk_{i}" for i in range(len(chunks) if "chunks" in dir() else 0)
            ]
            if chroma_ids:
                collection = get_or_create_collection(collection_name)
                collection.delete(ids=chroma_ids)
        except Exception:
            pass

        doc.ingestion_status = "failed"
        doc.ingestion_error = str(exc)
        await db.flush()

        logger.exception("Document %d ingestion failed: %s", document_id, exc)
        raise DocumentParseError(detail=f"文档解析失败: {exc}") from exc


async def remove_document_from_chroma(document_id: int, collection_name: str) -> None:
    """从 ChromaDB 中删除指定文档的所有 chunk。"""
    collection = get_or_create_collection(collection_name)
    try:
        # 按 metadata 过滤删除
        results = collection.get(
            where={"document_id": document_id},
            include=[],
        )
        if results and results["ids"]:
            collection.delete(ids=results["ids"])
            from backend.rag.retrieval import invalidate_cache
            invalidate_cache(collection_name)
            logger.info(
                "Removed %d chunks of document %d from collection %s",
                len(results["ids"]),
                document_id,
                collection_name,
            )
    except Exception as exc:
        logger.warning("Failed to remove doc %d from ChromaDB: %s", document_id, exc)


async def cleanup_kb_collection(knowledge_base_id: int, collection_name: str) -> None:
    """删除整个知识库的 ChromaDB collection。"""
    delete_collection(collection_name)
    logger.info("Cleaned up ChromaDB collection for KB %d: %s", knowledge_base_id, collection_name)
